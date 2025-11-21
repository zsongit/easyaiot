#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown转换为公众号可用的DOCX格式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_markdown_to_docx(md_content, output_path):
    """将Markdown转换为DOCX格式"""
    
    # 标题图标映射
    title_icons = {
        "引言": "💡",
        "项目概述": "📋",
        "设计哲学": "🎯",
        "平台定位": "📍",
        "核心价值": "✨",
        "技术架构": "🏗️",
        "模块化设计理念": "🧩",
        "数据流转架构": "🔄",
        "存储方案": "💾",
        "核心AI能力": "🤖",
        "全面的AI技术栈": "🔧",
        "革命性的零样本标注技术": "🚀",
        "多场景预训练模型": "📦",
        "IoT能力": "🌐",
        "完整的设备生命周期管理": "📱",
        "强大的规则引擎": "⚙️",
        "数据智能分析": "📊",
        "部署灵活性": "☁️",
        "独立部署优势": "🔀",
        "一键部署方案": "⚡",
        "应用场景适配": "🎨",
        "核心优势": "⭐",
        "多语言混编架构": "🔤",
        "零样本标注技术": "🎯",
        "云边端灵活部署": "🌍",
        "丰富生态支持": "🌳",
        "持续迭代优化": "🔄",
        "应用场景": "🎬",
        "人群密度管控": "👥",
        "周界防护": "🛡️",
        "跌倒检测": "⚠️",
        "异常逗留识别": "👀",
        "肢体冲突预警": "⚔️",
        "非法闯入检测": "🚫",
        "公共场所控烟": "🚭",
        "人流统计管控": "📈",
        "区域越界预警": "🚧",
        "环境安全检查": "🔍",
        "火灾早预警": "🔥",
        "扩展应用领域": "🔮",
        "系统展示": "🖼️",
        "核心功能界面展示": "💻",
        "技术实现": "💻",
        "设备控制核心逻辑": "🎮",
        "安全认证体系": "🔐",
        "AI模型管理": "🧠",
        "高性能任务处理": "⚡",
        "功能介绍": "📖",
        "设备管理模块": "📱",
        "流媒体管理模块": "🎥",
        "数据标注模块": "✏️",
        "模型训练与管理模块": "🎓",
        "AI智能分析模块": "🔬",
        "规则引擎模块": "⚙️",
        "系统管理模块": "🛠️",
        "数据统计与分析模块": "📊",
        "部署安装": "🚀",
        "部署要求": "📋",
        "部署优势": "✅",
        "社区与开源": "❤️",
        "我们的承诺": "🤝",
        "加入我们": "🌟",
        "演示环境与支持": "🌐",
        "在线演示": "💻",
        "结语": "🎉",
        "联系方式": "📞",
    }
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_language = ''
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # 处理代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line.strip()[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                # 添加代码块
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(40, 40, 40)
                # 设置段落背景色（通过表格实现）
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(24)
                run.font.bold = True
        elif line.startswith('## '):
            title = line[3:].strip()
            icon = get_icon_for_title(title, title_icons)
            p = doc.add_heading(f'{icon} {title}', level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(20)
                run.font.bold = True
        elif line.startswith('### '):
            title = line[4:].strip()
            icon = get_icon_for_title(title, title_icons)
            p = doc.add_heading(f'{icon} {title}', level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(18)
                run.font.bold = True
        elif line.startswith('#### '):
            title = line[5:].strip()
            icon = get_icon_for_title(title, title_icons)
            p = doc.add_heading(f'{icon} {title}', level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = '微软雅黑'
                run.font.size = Pt(16)
                run.font.bold = True
        # 处理引用
        elif line.startswith('> '):
            quote = line[2:].strip()
            p = doc.add_paragraph()
            p.style = 'Quote'
            add_formatted_text(p, quote)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(102, 102, 102)
        # 处理分隔线
        elif line.strip() == '---':
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(200, 200, 200)
        # 处理列表
        elif line.strip().startswith('- '):
            items = []
            while i < len(lines) and lines[i].strip().startswith('- '):
                item = lines[i][2:].strip()
                items.append(item)
                i += 1
            for item in items:
                p = doc.add_paragraph(item, style='List Bullet')
                add_formatted_text(p, item, is_new=False)
            continue
        elif re.match(r'^\d+\.\s', line):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i]):
                item = re.sub(r'^\d+\.\s', '', lines[i]).strip()
                items.append(item)
                i += 1
            for item in items:
                p = doc.add_paragraph(item, style='List Number')
                add_formatted_text(p, item, is_new=False)
            continue
        # 处理普通段落
        elif line.strip():
            content = line.strip()
            p = doc.add_paragraph()
            add_formatted_text(p, content)
        else:
            # 空行
            doc.add_paragraph()
        
        i += 1
    
    # 保存文档
    doc.save(output_path)

def get_icon_for_title(title, title_icons):
    """为标题获取合适的图标"""
    # 先尝试完整匹配
    if title in title_icons:
        return title_icons[title]
    
    # 尝试部分匹配
    for key, icon in title_icons.items():
        if key in title:
            return icon
    
    # 根据标题内容智能匹配
    if '模块' in title:
        return '📦'
    elif '管理' in title:
        return '🛠️'
    elif '技术' in title or '架构' in title:
        return '🏗️'
    elif '部署' in title or '安装' in title:
        return '🚀'
    elif '场景' in title or '应用' in title:
        return '🎬'
    elif '能力' in title or '功能' in title:
        return '✨'
    elif '分析' in title or '统计' in title:
        return '📊'
    elif '训练' in title or '模型' in title:
        return '🧠'
    elif '设备' in title:
        return '📱'
    elif '视频' in title or '流媒体' in title:
        return '🎥'
    elif '数据' in title:
        return '💾'
    elif '安全' in title:
        return '🔐'
    elif '规则' in title or '引擎' in title:
        return '⚙️'
    elif '演示' in title or '环境' in title:
        return '🌐'
    elif '社区' in title or '开源' in title:
        return '❤️'
    elif '联系' in title:
        return '📞'
    elif '结语' in title:
        return '🎉'
    else:
        return '📌'

def add_formatted_text(paragraph, text, is_new=True):
    """添加格式化文本（处理加粗、代码、链接）"""
    if is_new:
        paragraph.clear()
    
    # 处理加粗 **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 加粗文本
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.color.rgb = RGBColor(231, 76, 60)  # 红色加粗
        elif part.startswith('`') and part.endswith('`'):
            # 代码文本
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(52, 152, 219)
        else:
            # 普通文本
            # 处理链接 [text](url)
            link_parts = re.split(r'(\[[^\]]+\]\([^\)]+\))', part)
            for link_part in link_parts:
                match = re.match(r'\[([^\]]+)\]\(([^\)]+)\)', link_part)
                if match:
                    link_text = match.group(1)
                    link_url = match.group(2)
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(52, 152, 219)
                    run.underline = True
                    # 注意：python-docx不支持超链接，只能显示文本
                elif link_part:
                    paragraph.add_run(link_part)

if __name__ == '__main__':
    # 读取markdown文件
    with open('/projects/easyaiot/.doc/宣传文章/项目宣传2.md', 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为DOCX
    output_path = '/projects/easyaiot/.doc/宣传文章/doc/项目宣传2_公众号版.docx'
    convert_markdown_to_docx(md_content, output_path)
    
    print(f"转换完成！DOCX文件已保存到: {output_path}")
    print("您可以直接打开该文件，复制内容到公众号编辑器中使用。")
