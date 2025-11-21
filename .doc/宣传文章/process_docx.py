#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
处理docx文件，优化格式用于公众号导入
- 清理多余的换行
- 设置美观的字体
- 优化段落格式
- 智能合并短段落
"""

import re
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def is_title(text):
    """判断是否为标题"""
    # 标题特征：短文本、可能包含数字编号、可能全角标点结尾
    text = text.strip()
    if len(text) < 50 and (text.endswith('：') or text.endswith(':') or 
                          re.match(r'^[一二三四五六七八九十\d]+[、.．]', text) or
                          re.match(r'^第[一二三四五六七八九十\d]+[章节部分]', text)):
        return True
    return False

def should_merge(prev_text, current_text):
    """判断是否应该合并段落"""
    # 如果前一个段落很短（少于30字）且不是标题，且当前段落也不是标题，可以合并
    if len(prev_text) < 30 and not is_title(prev_text) and not is_title(current_text):
        # 如果前一个段落不以句号、问号、感叹号结尾，可以合并
        if not prev_text.endswith(('。', '！', '？', '.', '!', '?')):
            return True
    return False

def process_docx(input_path, output_path=None):
    """
    处理docx文件，优化格式
    
    Args:
        input_path: 输入docx文件路径
        output_path: 输出docx文件路径，如果为None则自动生成
    """
    if not os.path.exists(input_path):
        print(f"错误：文件不存在: {input_path}")
        return False
    
    # 如果没有指定输出路径，自动生成
    if output_path is None:
        base_name = os.path.splitext(input_path)[0]
        output_path = f"{base_name}_优化版.docx"
    
    # 读取文档
    doc = Document(input_path)
    
    # 公众号常用字体：微软雅黑、思源黑体、PingFang SC等
    font_name = "微软雅黑"  # 公众号常用字体
    
    # 收集所有段落文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 清理段落内的多余空格
            text = re.sub(r' +', ' ', text)
            # 清理段落内的多余换行（保留单个换行用于列表等）
            text = re.sub(r'\n{2,}', '\n', text)
            paragraphs.append(text)
    
    # 智能合并段落
    merged_paragraphs = []
    for i, text in enumerate(paragraphs):
        if not merged_paragraphs:
            merged_paragraphs.append(text)
        else:
            prev_text = merged_paragraphs[-1]
            # 如果应该合并，则合并
            if should_merge(prev_text, text):
                merged_paragraphs[-1] = prev_text + text
            else:
                merged_paragraphs.append(text)
    
    # 创建新文档
    new_doc = Document()
    
    # 设置默认字体
    style = new_doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(14)  # 公众号常用字号：14pt（更美观）
    
    # 添加处理后的段落
    for text in merged_paragraphs:
        # 如果段落中包含换行，可能是列表或需要保留的结构
        if '\n' in text:
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    para = new_doc.add_paragraph(line)
                    para_format = para.paragraph_format
                    para_format.space_after = Pt(6)  # 段落间距（更紧凑）
                    para_format.line_spacing = 1.5  # 行距1.5倍（更美观）
                    
                    # 判断是否为标题，设置相应格式
                    if is_title(line):
                        para_format.space_before = Pt(10)  # 标题前间距
                        para_format.space_after = Pt(8)  # 标题后间距
                        for run in para.runs:
                            run.font.name = font_name
                            run.font.size = Pt(16)  # 标题16pt（与正文14pt形成层次）
                            run.bold = True
                    else:
                        for run in para.runs:
                            run.font.name = font_name
                            run.font.size = Pt(14)  # 正文14pt（公众号常用）
        else:
            # 普通段落
            para = new_doc.add_paragraph(text)
            para_format = para.paragraph_format
            para_format.space_after = Pt(6)  # 段落间距（更紧凑）
            para_format.line_spacing = 1.5  # 行距1.5倍（更美观）
            
            # 判断是否为标题
            if is_title(text):
                para_format.space_before = Pt(10)  # 标题前间距
                para_format.space_after = Pt(8)  # 标题后间距
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(16)  # 标题16pt（与正文14pt形成层次）
                    run.bold = True
            else:
                for run in para.runs:
                    run.font.name = font_name
                    run.font.size = Pt(14)  # 正文14pt（公众号常用）
    
    # 保存新文档
    new_doc.save(output_path)
    print(f"处理完成！已保存到: {output_path}")
    print(f"原段落数: {len(paragraphs)}, 优化后段落数: {len(merged_paragraphs)}")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python3 process_docx.py <输入文件路径> [输出文件路径]")
        print("示例: python3 process_docx.py doc/项目宣传2_公众号版.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        process_docx(input_file, output_file)
    except ImportError:
        print("错误：需要安装python-docx库")
        print("请运行: pip install python-docx")
        sys.exit(1)
    except Exception as e:
        print(f"处理文件时出错: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

